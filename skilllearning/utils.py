import os
import time
from pathlib import Path
import pyautogui
from enum import Enum
from datetime import datetime, timedelta
import pygetwindow as gw

from docx import Document
from docx.shared import Inches
import cv2
import numpy as np

word = None
new_doc = None
doc = None
logon_window_title = ''
main_window_title = ''

img_sap_logon = "sap_logon.png"
img_huanyucsd = "hy_csd.png"
img_username = 'username.png'
img_password = 'password.png'
img_lang = 'lang.png'
img_active_work = 'active_work.png'
img_drag_down = 'drag_down.png'
img_input_tag = 'input_tag.png'
img_user_name = 'user_name.png'
img_job_status = ['released_check.png', 'ready_check.png', 'active_check.png', 'finished_check.png']
img_job_start_from = 'job_start_from.png'
img_execute = 'execute.png'
img_month = 'month.png'
img_day8 = 'day8.png'
img_list = 'list.png'
img_previous_hours = ['previous_hours_close.png', 'previous_hours_open.png']
img_snapshot = ['snapshot_close.png', 'snapshot_open.png']
img_serverid = ['serverid.png', 'serverid_day.png']
img_yestoday = 'yestoday.png'
img_nst02 = '/nst02'
img_close_yes = 'close_yes'
img_h_scroll = 'h_scroll.png'
pic_width_in_word = 5.8  # 图片在word里的宽度 inch

txt_username = 'basupport'
txt_password = 'sy1234'
txt_lang = 'En'
txt_nsm50 = '/nsm50'
txt_nsm37 = '/nsm37'
txt_ndb13 = '/ndb13'
txt_ndb12 = '/ndb12'
txt_nsm12 = '/nsm12'
txt_nst06 = '/nst06'
txt_nst03n = '/nst03n'
txt_nst22 = '/nst22'

img_offset = ["center"]

BASE_DIR = Path(__file__).resolve().parent.parent
Img_Path = os.path.join(BASE_DIR, 'skilllearning', 'imgs')
Screenshot_Path = os.path.join(BASE_DIR, 'skilllearning', 'screenshot')
DELAY = 1


# 定义一个枚举类
class Pos_Offset(Enum):
    LEFT = 1
    CENTER = 2
    RIGHT = 3


def get_timestamp():
    now = datetime.now()
    current_time_str = now.strftime('%Y%m%d%H%M%S')
    print(current_time_str)
    return current_time_str


def get_yestoday_text():
    # 获取当前日期
    current_date = datetime.now()
    # 计算前一天的日期
    yesterday = current_date - timedelta(days=1)
    # 将前一天的日期格式化为 'dd.mm.yyyy' 格式
    formatted_yesterday = yesterday.strftime('%d.%m.%Y')
    print(formatted_yesterday)
    return formatted_yesterday


def get_img_file(img_username):
    return os.path.join(Img_Path, img_username)


def get_screenshot_file(img_username):
    if not os.path.exists(Screenshot_Path):
        os.makedirs(Screenshot_Path, exist_ok=True)
    return os.path.join(Screenshot_Path, img_username)


def press_enter():
    pyautogui.press('enter')
    time.sleep(0.2)


def press_key(key_name):
    pyautogui.press(key_name)
    time.sleep(0.2)


def mouse_drag(x=0, y=0, duration=0.2):
    # 按住鼠标左键，用0.2秒钟把鼠标向上拖拽
    pyautogui.dragRel(x, y, duration=duration)


def get_cur_window_title():
    active_window = gw.getActiveWindow()
    print(active_window.title)
    return active_window.title


def close_window_by_title(title):
    # 获取所有标题包含"Notepad"的窗口
    windows = gw.getWindowsWithTitle(title)
    # 遍历窗口列表并关闭每个窗口
    for window in windows:
        window.close()


def find_icon(img_path, location: str = Pos_Offset.CENTER.name, pixs: int = 20):
    # 获得文件图片在现在的屏幕上面的坐标，返回的是一个元组(top, left, width, height)
    # 如果截图没找到，pyautogui.locateOnScreen()函数返回None
    a = ()
    try:
        a = pyautogui.locateOnScreen(img_path)
    except Exception as e:
        print("pyautogui.locateOnScreen 发生了一个错误：", str(e))

    if a:
        print(a)  # 打印结果为Box(left=0, top=0, width=300, height=400)
        if location == Pos_Offset.CENTER.name:
            x, y = pyautogui.center(a)  # 获得文件图片在现在的屏幕上面的中心坐标
            print(x, y)  # 打印结果为150 200
        elif location == Pos_Offset.RIGHT.name:
            x, y, width, height = a
            x = x + width + pixs
            y = y + int(height / 2)
            # x, y = pyautogui.locateCenterOnScreen(r'C:\Users\co_co\Desktop\PY\region_screenshot.png')  # 这步与上面的四行代码作用一样
        else:
            x, y, width, height = a
            x = x + pixs
            y = y + int(height / 2)
        time.sleep(0.2)
        return (x, y)
    else:
        print(f"没有找到图片：{img_path}")
        return (None, None)


def find_icon_click(img_path, location: str = Pos_Offset.CENTER.name, click: int = 1, pixs: int = 20):
    x, y = find_icon(img_path=img_path, location=location, pixs=pixs)
    if x:
        pyautogui.moveTo(x, y, duration=0.25)  # 移动到 (100,100)
        if click == 1:
            pyautogui.click()  # 鼠标当前位置点击一下
        elif click == 2:
            pyautogui.doubleClick()
        else:
            pyautogui.tripleClick()  # 鼠标当前位置左击三下
        time.sleep(DELAY)


def find_icon_input(img_path, location: str = Pos_Offset.CENTER.name, input_str: str = "a"):
    x, y = find_icon(img_path, location)
    if x:
        pyautogui.moveTo(x, y, duration=0.25)  # 移动到 (100,100)
        pyautogui.click()  # 鼠标当前位置点击一下
        # 全选文本框中的文本
        pyautogui.hotkey('ctrl', 'a')  # 或者在macOS上使用 pyautogui.hotkey('command', 'a')
        # 删除选中的文本
        pyautogui.press('backspace')  # 或者使用 pyautogui.press('delete')
        # 等待一小段时间，确保文本已经被清空
        time.sleep(0.5)
        print(input_str)
        pyautogui.typewrite(input_str, interval=0.25)  # 每次输入间隔0.25秒，输入Hello world!
    time.sleep(DELAY)


def find_icon_drag(img_path, location: str = Pos_Offset.CENTER.name, lx=0, ly=0, duration=0.2):
    x, y = find_icon(img_path, location)
    if x:
        pyautogui.moveTo(x, y, duration=0.25)  # 移动到 (100,100)
        pyautogui.dragRel(lx, ly, duration=duration)


def expand_window(ix: int = 50, iy: int = 50):
    active_window = gw.getActiveWindow()
    print(active_window.title)
    if active_window:
        # 获取窗口的位置和尺寸
        x, y, width, height = active_window.left, active_window.top, active_window.width, active_window.height
        lx = x + width - 3
        ly = y + height - 3
        pyautogui.moveTo(lx, ly)
        pyautogui.mouseDown()  # 鼠标左键按下再松开
        pyautogui.moveRel(ix, iy)
        pyautogui.mouseUp()


# 点击上周今天
def click_last_weekday(img_path, cur_day, last_day, type: int = 0):
    # 读取大图和小图
    pts = []
    if cur_day > 7:
        small_name = f'd_{last_day}.png'
    else:
        small_name = f'd_{last_day}.png'
        # img_path = ""
        x, y = find_icon(img_path=img_path, location=Pos_Offset.LEFT.name, pixs=3)
        pyautogui.moveTo(x, y, duration=0.25)  # 移动到 (100,100)
        for i in range(9):
            pyautogui.click()
            time.sleep(0.1)
    print("last_day-->", small_name)
    screen_name = action_screen_shot(20, 340, 480, 250, type)  # 648,182   220  #24,343  518,  500,250
    small_name = get_img_file(small_name)
    large_name = get_screenshot_file(screen_name)
    large_image = cv2.imread(large_name)
    small_image = cv2.imread(small_name)
    print("large-->", large_name)
    # 转换为灰度图像
    large_gray = cv2.cvtColor(large_image, cv2.COLOR_BGR2GRAY)
    small_gray = cv2.cvtColor(small_image, cv2.COLOR_BGR2GRAY)
    # 应用matchTemplate方法
    result = cv2.matchTemplate(large_gray, small_gray, cv2.TM_CCOEFF_NORMED)
    # 设置阈值
    threshold = 0.99
    # 找到所有匹配的位置
    loc = np.where(result >= threshold)
    active_window = gw.getActiveWindow()
    print(active_window.title)
    if active_window:
        # 获取窗口的位置和尺寸
        x, y, width, height = active_window.left, active_window.top, active_window.width, active_window.height
    # 绘制矩形
    for pt in zip(*loc[::-1]):
        cv2.rectangle(large_image, pt, (pt[0] + small_gray.shape[1], pt[1] + small_gray.shape[0]), (0, 255, 0), 2)
        print("pt-->", pt)
        pts.append((pt[0], pt[1]))
    max_pt = max(pts, key=lambda x: x[0])
    print(max_pt[0] + 20, max_pt[1] + 340)
    pyautogui.moveTo(max_pt[0] + 20 + x + 15, max_pt[1] + 340 + y + 10, duration=0.25)  # 移动到 (100,100)
    pyautogui.click()


# 截图
def action_screen_shot(off_x: int = 0, off_y: int = 0, off_w: int = 0, off_h: int = 0, type: int = 0):
    active_window = gw.getActiveWindow()
    print(active_window.title)
    if active_window:
        # 获取窗口的位置和尺寸
        x, y, width, height = active_window.left, active_window.top, active_window.width, active_window.height
        # 截图当前窗口
        if type == 0:
            screenshot = pyautogui.screenshot(region=(x + off_x, y + off_y, width + off_w, height + off_h))
        elif type == 1:
            screenshot = pyautogui.screenshot(region=(x + off_x, y + off_y, x + off_w, y + off_h))
        elif type == 2:
            screenshot = pyautogui.screenshot(region=(off_x, off_y, off_w, off_h))
        else:
            screenshot = pyautogui.screenshot(region=(x + off_x, y + off_y, x + off_w, y + off_h))
        time.sleep(0.2)
        # 保存截图
        screen_name = f'screen_{get_timestamp()}.png'
        screenshot.save(get_screenshot_file(screen_name))
        time.sleep(0.2)
        return screen_name
    else:
        print("No active window found.")


# 点击 图标
def action_click_icon(img, location, click: int = 1, pixs: int = 20, sleep: int = DELAY):
    img_path = get_img_file(img)
    print(img_path)
    find_icon_click(img_path=img_path, location=location, click=click, pixs=pixs)
    time.sleep(sleep)


# 粘贴截图 word
def action_paste_word(screen_name, doc):
    if screen_name:
        img_path = get_screenshot_file(screen_name)
        doc.add_picture(img_path, width=Inches(pic_width_in_word))  # 默认 5.8 inche
        # 插入一个回车符（即添加一个新段落）
        doc.add_paragraph()
        doc.add_paragraph()
        time.sleep(0.2)


# 输入文本
def action_input_text(img, location, input_str, sleep: int = DELAY):
    img_path = get_img_file(img)
    print(img_path)
    find_icon_input(img_path=img_path, location=location, input_str=input_str)
    press_enter()
    time.sleep(sleep)


def action_typewrite(text, interval: float = 0.25):
    pyautogui.hotkey('ctrl', 'a')  # 或者在macOS上使用 pyautogui.hotkey('command', 'a')
    # 删除选中的文本
    pyautogui.press('backspace')  # 或者使用 pyautogui.press('delete')
    # 等待一小段时间，确保文本已经被清空
    time.sleep(0.2)
    pyautogui.typewrite(text, interval=interval)


def action_click_arrow(img, off_x, off_y, sleep: int = DELAY):
    img_path = get_img_file(img[0])
    print(img_path)
    a = ()
    try:
        a = pyautogui.locateOnScreen(img_path)
        x, y, width, height = a
        pyautogui.moveTo(x + 8, y + 8, duration=0.25)  # 移动到 (100,100)
        pyautogui.click()  # 鼠标当前位置点击一下
        time.sleep(1)
        print("a-->", x, y)
    except Exception as e:
        print("pyautogui.locateOnScreen 发生了一个错误：", str(e))
    img_path = get_img_file(img[1])
    print(img_path)
    try:
        a = pyautogui.locateOnScreen(img_path)
        x, y, width, height = a
        pyautogui.moveTo(x + off_x, y + off_y, duration=0.25)  # 移动到 (100,100)
        pyautogui.click()  # 鼠标当前位置点击一下
        time.sleep(sleep)
        print("a-->", x, y)
    except Exception as e:
        print("pyautogui.locateOnScreen 发生了一个错误：", str(e))


# 展开树结构
def action_click_tree(img, off_x, off_y, sleep: int = DELAY):
    a = ()
    try:
        for i in img:
            img_path = get_img_file(i)
            a = pyautogui.locateOnScreen(img_path)
            x, y, width, height = a
            pyautogui.moveTo(x + 8, y + 8, duration=0.25)  # 移动到 (100,100)
            pyautogui.click()  # 鼠标当前位置点击一下
            time.sleep(1)
            print("a-->", x, y)

        pyautogui.moveTo(x + off_x, y + off_y, duration=0.25)
        pyautogui.doubleClick()
        time.sleep(sleep)
    except Exception as e:
        print("pyautogui.locateOnScreen 发生了一个错误：", str(e))


# 新建一个word文档，用今天日期命名，保存到桌面，并打开
def action_new_doc():
    global doc
    # 创建一个新的 Word 文档
    print("新建word文档")
    doc = Document()
    return


def action_save_doc(pre_name:str="sap"):
    global doc
    today_date = datetime.now().strftime('%Y%m%d%H%M%S')
    word_file_name = f'{pre_name}_{today_date}.docx'
    desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
    file_path = os.path.join(desktop_path, word_file_name)
    print("保存word文档-->", file_path)
    doc.save(file_path)


def action_crop_doc(x, y, w, h):
    #  截图，粘贴
    global doc
    print("截屏-->")
    screen_name = action_screen_shot(x, y, w, h, type=2)  # 648,182   220
    print("保存图片-->")
    action_paste_word(screen_name, doc)
    time.sleep(1)


def action_click_last_weekday():
    today = datetime.today()
    cur_day = today.day
    last_week_same_day = today - timedelta(weeks=1)
    last_day = last_week_same_day.day
    # img_path = ""
    img_name = get_img_file(img_h_scroll)
    print(img_name)
    click_last_weekday(img_name, cur_day, last_day, 1)
    time.sleep(1)
